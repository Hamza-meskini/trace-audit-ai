"""Synthetic dataset generator for TraceAudit evaluation and benchmarking.

Generates a realistic, synthetic technical document suite for an
Automotive Battery Control Unit (BCU) along with structured Ground Truth
JSON files covering 30 requirements across 5 verification scenarios:
  1. Supported (Verified)
  2. Partial Coverage
  3. Missing Evidence
  4. Potential Conflict
  5. Ambiguous / Needs Review

Outputs:
  - evaluation/documents/ (PDF, DOCX, XLSX files)
  - evaluation/ground_truth/ (requirements.json, evidence_links.json, expected_findings.json)
"""

import json
import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import fitz  # PyMuPDF


EVAL_DIR = Path(__file__).resolve().parent
DOCS_DIR = EVAL_DIR / "documents"
GT_DIR = EVAL_DIR / "ground_truth"


# ==============================================================================
# 30 GROUND TRUTH REQUIREMENTS DEFINITIONS
# ==============================================================================

BENCHMARK_REQUIREMENTS = [
    # ── Supported / Verified (10 requirements) ───────────────────────────────
    {
        "requirement_id": "REQ-BCU-001",
        "req_code": "REQ-BCU-001",
        "title": "High Voltage Pack Nominal Operating Range",
        "description": "The BCU shall monitor and regulate high-voltage traction packs operating from 400.0 V DC to 800.0 V DC under all standard driving and fast-charging conditions.",
        "category": "Electrical",
        "severity": "Critical",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 2, "relationship": "supports", "quote": "The BCU was tested across 400.0 V, 600.0 V and 800.0 V DC input levels with full telemetry accuracy."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-002",
        "req_code": "REQ-BCU-002",
        "title": "Continuous Discharge Current Monitoring",
        "description": "The primary current shunt sensing stage shall accurately measure continuous discharge currents up to 350.0 A with maximum error <= 0.5% full-scale.",
        "category": "Electrical",
        "severity": "High",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 2, "relationship": "supports", "quote": "Current sensor measured 350.0 A continuous discharge with 0.18% error across 4-hour soak test."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-003",
        "req_code": "REQ-BCU-003",
        "title": "Galvanic High-Voltage Isolation Barrier",
        "description": "The galvanic isolation barrier between high-voltage traction circuitry and 12V low-voltage logic domain shall withstand 2.5 kV AC RMS for 60 seconds with leakage current < 1.0 mA.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 3, "relationship": "supports", "quote": "Dielectric withstand test: 2,500 V AC applied for 60.0 s, measured leakage current 0.28 mA, verdict PASS."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-004",
        "req_code": "REQ-BCU-004",
        "title": "CAN-FD Telemetry Communication Rate",
        "description": "The powertrain telemetry interface shall support ISO 11898-2 CAN-FD protocol at 5.0 Mbps data phase rate with zero frame errors during full bus load.",
        "category": "Electrical",
        "severity": "Medium",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 3, "relationship": "supports", "quote": "CAN-FD bus evaluated at 5.0 Mbps across 10,000,000 frames with zero frame errors recorded."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-005",
        "req_code": "REQ-BCU-005",
        "title": "Quiescent Sleep State Current Draw",
        "description": "When vehicle ignition is off and BCU enters deep sleep mode, total quiescent current from the 12V auxiliary battery shall not exceed 150.0 uA.",
        "category": "Electrical",
        "severity": "Medium",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 2, "relationship": "supports", "quote": "Deep sleep quiescent current measured at 112.4 uA at 12.0 V auxiliary supply, well below 150.0 uA limit."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-006",
        "req_code": "REQ-BCU-006",
        "title": "Cell Voltage Sensing Measurement Precision",
        "description": "Individual battery cell voltage measurement accuracy shall be within +/- 1.5 mV across the entire 2.0 V to 4.5 V cell chemistry operating window.",
        "category": "Electrical",
        "severity": "High",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 2, "relationship": "supports", "quote": "Cell monitoring ASIC calibrated across 2.0V to 4.5V with maximum recorded error of +/- 0.85 mV."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-007",
        "req_code": "REQ-BCU-007",
        "title": "Safe State Contactor De-energization Latency",
        "description": "Upon detection of critical overvoltage or overtemperature fault, the safety supervisor shall open main DC contactors within maximum latency <= 10.0 ms.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "06_Thermal_Runaway_Safety_Report.pdf", "page": 2, "relationship": "supports", "quote": "Emergency contactor trip latency measured at 6.40 ms under simulated overvoltage event (threshold <= 10.0 ms)."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-008",
        "req_code": "REQ-BCU-008",
        "title": "Ingress Protection Rating IP67",
        "description": "The BCU aluminum die-cast housing and sealed automotive connectors shall satisfy IP67 ingress protection per ISO 20653 / IEC 60529.",
        "category": "Mechanical",
        "severity": "High",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "05_Environmental_EMC_Report.pdf", "page": 2, "relationship": "supports", "quote": "IP67 immersion test completed (1.0 meter submersion for 30 minutes). Zero water ingress detected inside enclosure."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-009",
        "req_code": "REQ-BCU-009",
        "title": "Pre-Charge Sequence Completion Time",
        "description": "The DC link pre-charge circuit sequence shall charge inverter DC-link capacitance to 95% of pack voltage within maximum duration <= 200.0 ms.",
        "category": "Electrical",
        "severity": "Medium",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 3, "relationship": "supports", "quote": "DC bus pre-charge timing test: 95% voltage reached in 142.0 ms with pre-charge resistor peak temperature within bounds."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-010",
        "req_code": "REQ-BCU-010",
        "title": "Solid-State Pyro-Fuse Trigger Circuit Latency",
        "description": "In hard short-circuit conditions, the solid-state pyro-fuse ignition circuit shall issue trigger pulse in <= 5.0 microseconds to isolate traction battery.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Supported",
        "expected_review_state": "Reviewed",
        "evidence_sources": [
            {"document": "06_Thermal_Runaway_Safety_Report.pdf", "page": 2, "relationship": "supports", "quote": "Hardware short-circuit analog comparator triggered pyro-switch in 2.8 microseconds."}
        ],
    },

    # ── Partial Coverage (6 requirements) ────────────────────────────────────
    {
        "requirement_id": "REQ-BCU-011",
        "req_code": "REQ-BCU-011",
        "title": "Extended Climatic Operating Temperature",
        "description": "The BCU shall maintain full functional performance and sensor accuracy across ambient operating temperature range from -40.0 °C to +85.0 °C.",
        "category": "Environmental",
        "severity": "High",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "05_Environmental_EMC_Report.pdf", "page": 2, "relationship": "partial_support", "quote": "Climatic chamber thermal testing performed from -20.0 °C to +70.0 °C under continuous telemetry logging."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-012",
        "req_code": "REQ-BCU-012",
        "title": "Damp Heat Cyclic Humidity Endurance",
        "description": "The controller shall withstand damp heat cyclic testing at 95.0% relative humidity and +65.0 °C for a continuous duration of 500.0 hours without insulation degradation.",
        "category": "Environmental",
        "severity": "Medium",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "05_Environmental_EMC_Report.pdf", "page": 2, "relationship": "partial_support", "quote": "Damp heat endurance test chamber run completed for 100.0 hours at 95% RH. Remainder of 500h test in progress."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-013",
        "req_code": "REQ-BCU-013",
        "title": "Thermal Runaway Gas Venting Pressure Calculation and Burst Validation",
        "description": "The pack pressure relief membrane must calculate thermal runaway gas release dynamics and demonstrate physical burst opening at pressure <= 300.0 mbar.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "06_Thermal_Runaway_Safety_Report.pdf", "page": 3, "relationship": "partial_support", "quote": "Computational fluid dynamics simulation and gas venting rate calculation completed for 300 mbar burst target; physical burst validation fixture pending."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-014",
        "req_code": "REQ-BCU-014",
        "title": "Tri-Axial Random Vibration Endurance Profile",
        "description": "The unit shall withstand random vibration per ISO 16750-3 across all three orthogonal axes (X, Y, and Z) for 32 hours per axis at 2.8 g RMS.",
        "category": "Mechanical",
        "severity": "High",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "05_Environmental_EMC_Report.pdf", "page": 3, "relationship": "partial_support", "quote": "Random vibration shaker table completed for X-axis and Y-axis at 2.8 g RMS. Z-axis evaluation awaiting test fixture adapter."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-015",
        "req_code": "REQ-BCU-015",
        "title": "Overcurrent Protection Multi-Tier Inverse Time Curve",
        "description": "Overcurrent trip response shall verify the complete multi-tier curve spanning 100.0 A to 500.0 A with programmable inverse-time characteristics.",
        "category": "Electrical",
        "severity": "High",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "04_Battery_Management_Test_Report.pdf", "page": 3, "relationship": "partial_support", "quote": "Single-point overcurrent verification conducted at 200.0 A setting only; multi-point curve characterization incomplete."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-016",
        "req_code": "REQ-BCU-016",
        "title": "Contact Resistance Degradation After Power Cycling",
        "description": "Terminal power lug contact resistance shall remain <= 0.20 mOhm after 100,000 power cycles as verified by micro-ohmmeter measurement record.",
        "category": "Electrical",
        "severity": "Medium",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "07_Compliance_Verification_Matrix.xlsx", "page": 1, "relationship": "partial_support", "quote": "Requirement ID: REQ-BCU-016; Description: Contact Resistance; Status: In Progress 25,000 cycles completed (0.11 mOhm)."}
        ],
    },

    # ── Missing Evidence (7 requirements) ────────────────────────────────────
    {
        "requirement_id": "REQ-BCU-017",
        "req_code": "REQ-BCU-017",
        "title": "ISO 21434 Hardware Root-of-Trust Secure Boot Signature",
        "description": "The BCU microcontroller shall enforce secure boot verification with ECDSA P-384 hardware root-of-trust key validation to prevent unauthorized firmware execution.",
        "category": "Cybersecurity",
        "severity": "Critical",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-018",
        "req_code": "REQ-BCU-018",
        "title": "High Voltage Interlock Loop (HVIL) Fast Disconnect Latency",
        "description": "Any break or interruption in the continuous high-voltage interlock loop shall trigger active discharge and contactor opening in <= 5.0 ms.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-019",
        "req_code": "REQ-BCU-019",
        "title": "High Altitude Barometric Pressure Operation",
        "description": "The BCU shall operate without high-voltage corona discharge or thermal overheating at altitudes up to 4000.0 m (61.6 kPa atmospheric pressure).",
        "category": "Environmental",
        "severity": "Medium",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-020",
        "req_code": "REQ-BCU-020",
        "title": "Salt Spray Atmosphere Corrosion Resistance",
        "description": "External connectors and housing shall show no galvanic corrosion or pitting after 720.0 hours neutral salt spray test per ASTM B117 / ISO 9227.",
        "category": "Mechanical",
        "severity": "Medium",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-021",
        "req_code": "REQ-BCU-021",
        "title": "Radiated RF Electromagnetic Immunity 100 V/m",
        "description": "The BCU shall maintain uninterrupted operation during radiated RF field immunity exposure of 100.0 V/m from 200.0 MHz to 2.0 GHz per ISO 11452-2.",
        "category": "Electrical",
        "severity": "High",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-022",
        "req_code": "REQ-BCU-022",
        "title": "Mean Time Between Failures Reliability MTBF Target",
        "description": "Theoretical system reliability MTBF shall exceed 250,000 hours continuous operation at 40°C ambient calculated per Telcordia SR-332 / SN 29500.",
        "category": "Documentation",
        "severity": "Medium",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },
    {
        "requirement_id": "REQ-BCU-023",
        "req_code": "REQ-BCU-023",
        "title": "State of Charge (SOC) Extended Kalman Filter Estimation Accuracy",
        "description": "The battery state-of-charge estimator shall maintain SOC tracking drift < 2.0% across dynamic WLTP and US06 drive cycles over 1,000 operating hours.",
        "category": "Documentation",
        "severity": "High",
        "expected_status": "Missing",
        "expected_review_state": "Open",
        "evidence_sources": [],
    },

    # ── Potential Conflict (5 requirements) ──────────────────────────────────
    {
        "requirement_id": "REQ-BCU-024",
        "req_code": "REQ-BCU-024",
        "title": "Upper Limit Traction Pack Voltage Rating",
        "description": "The BCU sensing and isolation design shall operate across 400.0 V DC to 800.0 V DC traction bus voltage envelope.",
        "category": "Electrical",
        "severity": "Critical",
        "expected_status": "Conflict",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "01_Product_Requirements_SRS.docx", "page": None, "relationship": "specification", "quote": "The system must remain fully functional across an input tolerance range of 400.0 V DC to 800.0 V DC."},
            {"document": "03_OEM_Supplier_Datasheet.pdf", "page": 1, "relationship": "contradicts", "quote": "Maximum rated pack sensing voltage: 400.0 V to 750.0 V DC maximum. Operation above 750V is not supported."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-025",
        "req_code": "REQ-BCU-025",
        "title": "Maximum Ambient Operating Thermal Limit",
        "description": "The power electronics stage shall operate up to +85.0 °C ambient temperature under continuous current load.",
        "category": "Environmental",
        "severity": "High",
        "expected_status": "Conflict",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "01_Product_Requirements_SRS.docx", "page": None, "relationship": "specification", "quote": "The equipment shall maintain full functional performance across an ambient operating temperature range of -40.0 °C to +85.0 °C."},
            {"document": "03_OEM_Supplier_Datasheet.pdf", "page": 1, "relationship": "contradicts", "quote": "Operating Ambient Temperature: -10.0 °C to +70.0 °C. Operation above +70°C causes thermal shutdown."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-026",
        "req_code": "REQ-BCU-026",
        "title": "Diagnostic Service Port Access Authentication",
        "description": "Unified Diagnostic Services (UDS) over Ethernet shall mandate cryptographic seed-key authentication before permitting flashing or calibration access.",
        "category": "Cybersecurity",
        "severity": "High",
        "expected_status": "Conflict",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "01_Product_Requirements_SRS.docx", "page": None, "relationship": "specification", "quote": "UDS service interface requires authenticated credential login and cryptographic seed-key security."},
            {"document": "02_System_Architecture_Spec.pdf", "page": 2, "relationship": "contradicts", "quote": "Engineering service port provides open access with no login required for factory automated calibration."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-027",
        "req_code": "REQ-BCU-027",
        "title": "Galvanic Isolation vs Ground Architecture",
        "description": "The high-voltage analog measurement sub-circuit shall maintain complete galvanic isolation from the vehicle chassis ground.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Conflict",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "01_Product_Requirements_SRS.docx", "page": None, "relationship": "specification", "quote": "All primary sensing channels shall feature galvanic optical and magnetic isolation from vehicle ground."},
            {"document": "03_OEM_Supplier_Datasheet.pdf", "page": 1, "relationship": "contradicts", "quote": "Architecture Note: Low-side sensing circuitry uses a non-isolated shared ground tied directly to chassis return."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-028",
        "req_code": "REQ-BCU-028",
        "title": "Active Cell Balancing Discharge Current",
        "description": "Active cell balancing circuit shall deliver balancing current of 200.0 mA per channel to equalize pack state-of-charge.",
        "category": "Electrical",
        "severity": "Medium",
        "expected_status": "Conflict",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "01_Product_Requirements_SRS.docx", "page": None, "relationship": "specification", "quote": "Cell balancing circuit operates between 150 mA and 200.0 mA discharge current."},
            {"document": "03_OEM_Supplier_Datasheet.pdf", "page": 1, "relationship": "contradicts", "quote": "Integrated balancing switch operates between 50 mA and 100.0 mA maximum continuous dissipation."}
        ],
    },

    # ── Ambiguous / Human Review Cases (2 requirements) ──────────────────────
    {
        "requirement_id": "REQ-BCU-029",
        "req_code": "REQ-BCU-029",
        "title": "Electrolyte Leakage Detection Acoustic Impedance Threshold",
        "description": "The ultrasonic acoustic sensing channel shall flag electrolyte leakage when acoustic signature dampening exceeds 12.0 dB relative to dry baseline.",
        "category": "Safety",
        "severity": "High",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "06_Thermal_Runaway_Safety_Report.pdf", "page": 3, "relationship": "ambiguous", "quote": "Acoustic impedance sensor characterization demonstrated signal dampening in presence of liquid solvent, but exact dB threshold calibration requires vehicle pack installation test."}
        ],
    },
    {
        "requirement_id": "REQ-BCU-030",
        "req_code": "REQ-BCU-030",
        "title": "High Voltage Pyro-Switch Minimum Firing Energy",
        "description": "Capacitive firing module shall store minimum 4.5 J energy at 24V booster stage to guarantee single-pulse detonation of the main safety disconnector.",
        "category": "Safety",
        "severity": "Critical",
        "expected_status": "Partial",
        "expected_review_state": "Needs review",
        "evidence_sources": [
            {"document": "06_Thermal_Runaway_Safety_Report.pdf", "page": 2, "relationship": "partial_support", "quote": "Firing capacitor energy calculation confirmed 4.8 J nominal theoretical capacity; physical discharge waveform capture pending final harness qualification."}
        ],
    },
]


# ==============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ==============================================================================

def create_product_requirements_srs():
    """Create 01_Product_Requirements_SRS.docx with all 30 requirements."""
    doc = Document()

    title = doc.add_heading("System Requirements Specification (SRS)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Product: Automotive Battery Control Unit (BCU-800V)\n").bold = True
    p.add_run("Document Ref: SRS-AUT-BCU-800V | Revision: 3.2 | Status: Released\n")
    p.add_run("System Architecture: 800V Silicon-Carbide EV Platform\n")
    p.add_run("Applicable Standards: ISO 26262 (ASIL D), ISO 6469-3, UN ECE R100, ISO 21434, IEC 60529\n")

    doc.add_heading("1. Purpose and System Architecture Overview", level=1)
    doc.add_paragraph(
        "The BCU-800V is an automotive safety-critical Electronic Control Unit (ECU) responsible for "
        "monitoring, state estimation, contactor sequencing, thermal supervision, and high-voltage isolation "
        "of the 800V traction battery system. This specification defines all mandatory electrical, safety, "
        "environmental, and cybersecurity requirements for vehicle type approval."
    )

    doc.add_heading("2. Technical Requirements Register", level=1)

    for req in BENCHMARK_REQUIREMENTS:
        p_req = doc.add_paragraph()
        p_req.add_run(f"{req['req_code']}: {req['title']}\n").bold = True
        p_req.add_run(f"Severity: {req['severity']} | Category: {req['category']}\n")
        p_req.add_run(f"{req['description']}\n")

    doc_path = DOCS_DIR / "01_Product_Requirements_SRS.docx"
    doc.save(str(doc_path))
    print(f"Generated: {doc_path}")


def create_system_architecture_pdf():
    """Create 02_System_Architecture_Spec.pdf with architectural diagrams and security conflict."""
    pdf = fitz.open()

    page1 = pdf.new_page(width=595, height=842)
    page1.insert_text((50, 50), "VOLT-DRIVE AUTOMOTIVE — SYSTEM ARCHITECTURE SPECIFICATION", fontsize=10, fontname="hebo", color=(0.1, 0.3, 0.6))
    page1.draw_line((50, 60), (545, 60), color=(0.1, 0.3, 0.6), width=1.5)

    page1.insert_text((50, 95), "BCU-800V System Architecture & Interface Design", fontsize=16, fontname="hebo")
    page1.insert_text((50, 115), "Document Ref: ARCH-BCU-2026-V1.0 | ASIL D Safety Architecture", fontsize=10, fontname="helv")

    page1.insert_text((50, 150), "1. High Voltage Domain Architecture & Sizing", fontsize=13, fontname="hebo")
    arch_desc = (
        "The traction pack monitoring architecture interfaces directly with 192 series-connected lithium-ion cells "
        "in an 800V nominal configuration (operating range 400.0 V to 800.0 V DC). Primary current measurement is performed "
        "via a 350.0 A continuous shunt sensor on the negative bus. The high-voltage domain is galvanically separated "
        "from the vehicle chassis and 12V body ground by a 2.5 kV AC reinforced insulation barrier."
    )
    page1.insert_textbox(fitz.Rect(50, 170, 545, 270), arch_desc, fontsize=9.5, fontname="helv")

    page1.insert_text((50, 290), "2. Service Port and Calibration Interface [REQ-BCU-026]", fontsize=13, fontname="hebo")
    sec_desc = (
        "Factory Calibration and Diagnostic Port: For high-speed end-of-line flashing and bench testing, the secondary "
        "Ethernet diagnostic interface is configured for factory deployment. Engineering service port provides open access "
        "with no login required for factory automated calibration and flashing routines. "
        "Production runtime builds will disable this port via OTP e-fuse configuration."
    )
    page1.insert_textbox(fitz.Rect(50, 310, 545, 410), sec_desc, fontsize=9.5, fontname="helv")

    pdf_path = DOCS_DIR / "02_System_Architecture_Spec.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_oem_supplier_datasheet_pdf():
    """Create 03_OEM_Supplier_Datasheet.pdf containing parameter contradictions."""
    pdf = fitz.open()

    page = pdf.new_page(width=595, height=842)
    page.insert_text((50, 50), "NEXUS POWER SEMICONDUCTOR — OEM ASIC COMPONENT DATASHEET", fontsize=10, fontname="hebo", color=(0.8, 0.2, 0.1))
    page.draw_line((50, 60), (545, 60), color=(0.8, 0.2, 0.1), width=1.5)

    page.insert_text((50, 95), "NX-BMS-800 Multi-Cell Monitoring & Balancing Front-End", fontsize=16, fontname="hebo")
    page.insert_text((50, 115), "Component Specification for Automotive High Voltage BMS", fontsize=10, fontname="helv")

    page.insert_text((50, 150), "1. Absolute Maximum Ratings & Operating Limits", fontsize=12, fontname="hebo")

    # Table with deliberate contradictions
    page.draw_rect(fitz.Rect(50, 170, 545, 360), color=(0.7, 0.7, 0.7), width=1)
    page.draw_line((50, 195), (545, 195), color=(0.7, 0.7, 0.7), width=1)
    page.insert_text((60, 187), "Parameter", fontsize=10, fontname="hebo")
    page.insert_text((240, 187), "Datasheet Limit", fontsize=10, fontname="hebo")
    page.insert_text((390, 187), "Engineering Notes", fontsize=10, fontname="hebo")

    rows = [
        ("Pack Voltage Sensing Range", "400.0 V to 750.0 V DC", "Maximum rated pack sensing voltage: 400.0 V to 750.0 V DC maximum. Operation above 750V is not supported."),
        ("Operating Ambient Temperature", "-10.0 °C to +70.0 °C", "Operating Ambient Temperature: -10.0 °C to +70.0 °C. Operation above +70°C causes thermal shutdown."),
        ("Cell Balancing Current", "50 mA to 100.0 mA max", "Integrated balancing switch operates between 50 mA and 100.0 mA maximum continuous dissipation."),
        ("Ground Reference Architecture", "Shared common ground", "Architecture Note: Low-side sensing circuitry uses a non-isolated shared ground tied directly to chassis return."),
    ]
    y = 215
    for p_name, p_val, notes in rows:
        page.insert_text((60, y), p_name, fontsize=8.5, fontname="helv")
        page.insert_text((240, y), p_val, fontsize=8.5, fontname="hebo")
        page.insert_textbox(fitz.Rect(390, y - 10, 540, y + 25), notes, fontsize=7.5, fontname="helv")
        y += 35

    pdf_path = DOCS_DIR / "03_OEM_Supplier_Datasheet.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_battery_management_test_report_pdf():
    """Create 04_Battery_Management_Test_Report.pdf multi-page lab report."""
    pdf = fitz.open()

    # Page 1: Overview
    p1 = pdf.new_page(width=595, height=842)
    p1.insert_text((50, 50), "APEX AUTOMOTIVE CERTIFICATION LABS — CONFORMITY TEST REPORT", fontsize=10, fontname="hebo", color=(0.1, 0.5, 0.3))
    p1.draw_line((50, 60), (545, 60), color=(0.1, 0.5, 0.3), width=1.5)

    p1.insert_text((50, 95), "TEST REPORT: TR-BMS-2026-8801", fontsize=16, fontname="hebo")
    p1.insert_text((50, 115), "EUT: Automotive Battery Control Unit BCU-800V | HW Rev 3.0", fontsize=10, fontname="helv")
    p1.insert_text((50, 130), "Standards: ISO 26262-5:2018, UN ECE R100 Rev 3, ISO 6469-3", fontsize=10, fontname="helv")

    p1.insert_text((50, 160), "1. Executive Test Summary", fontsize=12, fontname="hebo")
    summary = (
        "This official laboratory test report certifies electrical compliance and safety testing of the BCU-800V. "
        "Testing was conducted under ISO/IEC 17025 accredited laboratory conditions across input voltage, current sensing, "
        "isolation breakdown, sleep power, and cell monitoring precision."
    )
    p1.insert_textbox(fitz.Rect(50, 175, 545, 235), summary, fontsize=9.5, fontname="helv")

    # Page 2: Electrical Verification Data
    p2 = pdf.new_page(width=595, height=842)
    p2.insert_text((50, 50), "APEX CERTIFICATION LABS | Report TR-BMS-2026-8801 | Page 2 of 3", fontsize=9, fontname="helv")
    p2.insert_text((50, 85), "2. Electrical Operating Range and Measurement Precision", fontsize=13, fontname="hebo")

    sec2 = (
        "Pack Voltage Sweep (REQ-BCU-001):\n"
        "The BCU was tested across 400.0 V, 600.0 V and 800.0 V DC input levels with full telemetry accuracy. "
        "Zero overvoltage latching or false alarms recorded across 50 power cycling iterations.\n\n"
        "Current Shunt Sensor Accuracy (REQ-BCU-002):\n"
        "Current sensor measured 350.0 A continuous discharge with 0.18% error across 4-hour soak test. "
        "Thermal rise at shunt terminals was 18.2 °C above ambient.\n\n"
        "Cell Voltage Measurement Precision (REQ-BCU-006):\n"
        "Cell monitoring ASIC calibrated across 2.0V to 4.5V with maximum recorded error of +/- 0.85 mV. "
        "Total cell sensing channel dispersion satisfied the +/- 1.5 mV tolerance requirement.\n\n"
        "Quiescent Sleep Current (REQ-BCU-005):\n"
        "Deep sleep quiescent current measured at 112.4 uA at 12.0 V auxiliary supply, well below 150.0 uA limit."
    )
    p2.insert_textbox(fitz.Rect(50, 105, 545, 380), sec2, fontsize=9.5, fontname="helv")

    # Page 3: High Voltage Isolation and Fast Sequencing
    p3 = pdf.new_page(width=595, height=842)
    p3.insert_text((50, 50), "APEX CERTIFICATION LABS | Report TR-BMS-2026-8801 | Page 3 of 3", fontsize=9, fontname="helv")
    p3.insert_text((50, 85), "3. Dielectric Isolation, Pre-Charge & Telemetry Verification", fontsize=13, fontname="hebo")

    sec3 = (
        "Dielectric Withstand Testing (REQ-BCU-003):\n"
        "Dielectric withstand test: 2,500 V AC applied for 60.0 s, measured leakage current 0.28 mA, verdict PASS. "
        "Test conducted between HV pack bus and low-voltage 12V harness.\n\n"
        "Pre-Charge Circuit Timing (REQ-BCU-009):\n"
        "DC bus pre-charge timing test: 95% voltage reached in 142.0 ms with pre-charge resistor peak temperature within bounds.\n\n"
        "CAN-FD Telemetry Verification (REQ-BCU-004):\n"
        "CAN-FD bus evaluated at 5.0 Mbps across 10,000,000 frames with zero frame errors recorded under 90% bus utilization.\n\n"
        "Overcurrent Verification Note (REQ-BCU-015):\n"
        "Single-point overcurrent verification conducted at 200.0 A setting only; multi-point curve characterization incomplete."
    )
    p3.insert_textbox(fitz.Rect(50, 105, 545, 390), sec3, fontsize=9.5, fontname="helv")

    pdf_path = DOCS_DIR / "04_Battery_Management_Test_Report.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_environmental_emc_report_pdf():
    """Create 05_Environmental_EMC_Report.pdf with IP67 pass and partial thermal/vibration."""
    pdf = fitz.open()

    p1 = pdf.new_page(width=595, height=842)
    p1.insert_text((50, 50), "DEKRA INDUSTRIAL LABS — ENVIRONMENTAL & EMC TEST REPORT", fontsize=10, fontname="hebo", color=(0.2, 0.4, 0.7))
    p1.draw_line((50, 60), (545, 60), color=(0.2, 0.4, 0.7), width=1.5)

    p1.insert_text((50, 95), "TEST REPORT: TR-ENV-2026-4412", fontsize=16, fontname="hebo")
    p1.insert_text((50, 115), "Product: BCU-800V Enclosure & Assembly | Test Standard: ISO 16750 / ISO 20653", fontsize=10, fontname="helv")

    sec1 = (
        "Ingress Protection Verification (REQ-BCU-008):\n"
        "IP67 immersion test completed (1.0 meter submersion for 30 minutes). Zero water ingress detected inside enclosure. "
        "Dust test IP6X completed with continuous depression of 2 kPa for 8 hours without dust penetration. Verdict: PASS IP67.\n\n"
        "Thermal Chamber Operating Range Test (REQ-BCU-011):\n"
        "Climatic chamber thermal testing performed from -20.0 °C to +70.0 °C under continuous telemetry logging. "
        "Operation at extremes of -40.0°C and +85.0°C not tested during this campaign due to chamber temperature limitation.\n\n"
        "Damp Heat Humidity Testing (REQ-BCU-012):\n"
        "Damp heat endurance test chamber run completed for 100.0 hours at 95% RH. Remainder of 500h test in progress.\n\n"
        "Mechanical Vibration Shaker Table (REQ-BCU-014):\n"
        "Random vibration shaker table completed for X-axis and Y-axis at 2.8 g RMS. Z-axis evaluation awaiting test fixture adapter."
    )
    p1.insert_textbox(fitz.Rect(50, 140, 545, 450), sec1, fontsize=9.5, fontname="helv")

    pdf_path = DOCS_DIR / "05_Environmental_EMC_Report.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_thermal_runaway_safety_report_pdf():
    """Create 06_Thermal_Runaway_Safety_Report.pdf covering functional safety and partials."""
    pdf = fitz.open()

    p1 = pdf.new_page(width=595, height=842)
    p1.insert_text((50, 50), "TUV RHEINLAND — FUNCTIONAL SAFETY VALIDATION REPORT", fontsize=10, fontname="hebo", color=(0.7, 0.2, 0.2))
    p1.draw_line((50, 60), (545, 60), color=(0.7, 0.2, 0.2), width=1.5)

    p1.insert_text((50, 95), "SAFETY REPORT: TR-SAF-2026-902", fontsize=16, fontname="hebo")
    p1.insert_text((50, 115), "Product: BCU-800V Safe State & Pyro-Switch Subsystem | ISO 26262 ASIL D", fontsize=10, fontname="helv")

    sec1 = (
        "1. Contactor De-Energization Latency (REQ-BCU-007):\n"
        "Emergency contactor trip latency measured at 6.40 ms under simulated overvoltage event (threshold <= 10.0 ms). "
        "Dual-channel high-side and low-side driver de-energized simultaneously without contact welding.\n\n"
        "2. Pyro-Fuse Ignition Timing (REQ-BCU-010):\n"
        "Hardware short-circuit analog comparator triggered pyro-switch in 2.8 microseconds. "
        "Total fault current interruption accomplished in 1.1 ms.\n\n"
        "3. Pressure Relief & Gas Venting Dynamics (REQ-BCU-013):\n"
        "Computational fluid dynamics simulation and gas venting rate calculation completed for 300 mbar burst target; "
        "physical burst validation fixture pending.\n\n"
        "4. Acoustic Leakage Impedance Analysis (REQ-BCU-029):\n"
        "Acoustic impedance sensor characterization demonstrated signal dampening in presence of liquid solvent, "
        "but exact dB threshold calibration requires vehicle pack installation test.\n\n"
        "5. Firing Circuit Capacitive Energy (REQ-BCU-030):\n"
        "Firing capacitor energy calculation confirmed 4.8 J nominal theoretical capacity; physical discharge waveform capture pending final harness qualification."
    )
    p1.insert_textbox(fitz.Rect(50, 140, 545, 520), sec1, fontsize=9.5, fontname="helv")

    pdf_path = DOCS_DIR / "06_Thermal_Runaway_Safety_Report.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_compliance_verification_matrix_xlsx():
    """Create 07_Compliance_Verification_Matrix.xlsx summarizing compliance status."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "BCU_Compliance_Matrix"

    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    header_font = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
    data_font = Font(name="Segoe UI", size=9)

    headers = [
        "Requirement ID",
        "Requirement Title",
        "Target Value / Unit",
        "Test Report Ref",
        "Test Status",
        "Coverage Verdict",
        "Remarks / Engineering Notes",
    ]
    ws.append(headers)

    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    matrix_rows = [
        ("REQ-BCU-001", "Nominal Pack Voltage Range", "400.0 V to 800.0 V DC", "TR-BMS-2026-8801", "Completed", "PASS", "Fully tested across 400V, 600V, 800V"),
        ("REQ-BCU-002", "Continuous Discharge Current", "350.0 A max", "TR-BMS-2026-8801", "Completed", "PASS", "0.18% error measured at 350A"),
        ("REQ-BCU-003", "Dielectric Withstand Voltage", "2,500 V AC RMS", "TR-BMS-2026-8801", "Completed", "PASS", "Leakage 0.28 mA (limit < 1.0 mA)"),
        ("REQ-BCU-004", "CAN-FD Telemetry Rate", "5.0 Mbps", "TR-BMS-2026-8801", "Completed", "PASS", "Zero frame errors over 10M frames"),
        ("REQ-BCU-005", "Sleep Quiescent Current", "<= 150.0 uA", "TR-BMS-2026-8801", "Completed", "PASS", "112.4 uA measured"),
        ("REQ-BCU-006", "Cell Voltage Precision", "+/- 1.5 mV", "TR-BMS-2026-8801", "Completed", "PASS", "+/- 0.85 mV max error recorded"),
        ("REQ-BCU-007", "Contactor Trip Latency", "<= 10.0 ms", "TR-SAF-2026-902", "Completed", "PASS", "6.40 ms measured trip latency"),
        ("REQ-BCU-008", "Ingress Protection", "IP67", "TR-ENV-2026-4412", "Completed", "PASS", "Zero water penetration in 1m immersion"),
        ("REQ-BCU-009", "Pre-Charge Time", "<= 200.0 ms", "TR-BMS-2026-8801", "Completed", "PASS", "142.0 ms measured precharge time"),
        ("REQ-BCU-010", "Pyro-Switch Trigger Latency", "<= 5.0 us", "TR-SAF-2026-902", "Completed", "PASS", "2.8 us analog comparator response"),
        ("REQ-BCU-011", "Operating Temperature", "-40°C to +85°C", "TR-ENV-2026-4412", "Partial", "PARTIAL", "Tested -20°C to +70°C only [GAP]"),
        ("REQ-BCU-012", "Damp Heat Humidity", "95% RH for 500h", "TR-ENV-2026-4412", "Partial", "PARTIAL", "100h completed out of 500h [GAP]"),
        ("REQ-BCU-013", "Gas Venting Dynamics", "<= 300 mbar burst", "TR-SAF-2026-902", "Partial", "PARTIAL", "Simulation only; physical burst pending"),
        ("REQ-BCU-014", "Vibration Profile", "ISO 16750-3 Tri-axial", "TR-ENV-2026-4412", "Partial", "PARTIAL", "X/Y tested; Z axis pending"),
        ("REQ-BCU-015", "Overcurrent Curve", "100A to 500A curve", "TR-BMS-2026-8801", "Partial", "PARTIAL", "200A single point verified only"),
        ("REQ-BCU-016", "Contact Resistance", "<= 0.20 mOhm", "TR-MEC-2026-002", "In Progress", "PARTIAL", "25,000 cycles completed (0.11 mOhm)"),
        ("REQ-BCU-017", "Secure Boot Root of Trust", "ECDSA P-384", "NONE", "Not Started", "MISSING", "No test report or certificate in file"),
        ("REQ-BCU-018", "HVIL Response Time", "<= 5.0 ms", "NONE", "Not Started", "MISSING", "HVIL dynamic test record missing"),
        ("REQ-BCU-019", "Altitude 4000m Operation", "61.6 kPa", "NONE", "Not Started", "MISSING", "Altitude chamber report missing"),
        ("REQ-BCU-020", "Salt Spray Corrosion", "720h ASTM B117", "NONE", "Not Started", "MISSING", "Corrosion report missing"),
        ("REQ-BCU-021", "Radiated RF Immunity", "100 V/m", "NONE", "Not Started", "MISSING", "EMC radiated immunity test missing"),
        ("REQ-BCU-022", "Reliability MTBF", ">= 250,000 hours", "NONE", "Not Started", "MISSING", "Telcordia reliability calculation missing"),
        ("REQ-BCU-023", "SOC Kalman Filter Drift", "< 2.0% over 1000h", "NONE", "Not Started", "MISSING", "Drive cycle tracking validation missing"),
        ("REQ-BCU-024", "Upper Pack Voltage Rating", "800.0 V DC", "DS-NX-800", "Contradiction", "CONFLICT", "Supplier Datasheet limits max voltage to 750.0 V DC"),
        ("REQ-BCU-025", "Max Ambient Thermal Limit", "+85.0 °C", "DS-NX-800", "Contradiction", "CONFLICT", "Supplier Datasheet derates/limits at +70.0 °C"),
        ("REQ-BCU-026", "Diagnostic Authentication", "Seed-Key Credential", "ARCH-BCU-2026", "Contradiction", "CONFLICT", "System spec describes unauthenticated open access"),
        ("REQ-BCU-027", "Galvanic Isolation vs Ground", "Galvanic Isolation", "DS-NX-800", "Contradiction", "CONFLICT", "Datasheet specifies non-isolated shared ground"),
        ("REQ-BCU-028", "Cell Balancing Current", "200.0 mA", "DS-NX-800", "Contradiction", "CONFLICT", "Datasheet limits balancing to 100.0 mA"),
        ("REQ-BCU-029", "Acoustic Leak Dampening", "12.0 dB threshold", "TR-SAF-2026-902", "Ambiguous", "PARTIAL", "Vehicle installation validation required"),
        ("REQ-BCU-030", "Pyro Firing Energy", ">= 4.5 J", "TR-SAF-2026-902", "Partial", "PARTIAL", "Calculation completed; physical pulse pending"),
    ]

    for row_data in matrix_rows:
        ws.append(row_data)

    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 16
    ws.column_dimensions["G"].width = 40

    xlsx_path = DOCS_DIR / "07_Compliance_Verification_Matrix.xlsx"
    wb.save(str(xlsx_path))
    print(f"Generated: {xlsx_path}")


# ==============================================================================
# GROUND TRUTH EXPORT FUNCTIONS
# ==============================================================================

def export_ground_truth_json():
    """Export requirements.json, evidence_links.json, and expected_findings.json."""
    GT_DIR.mkdir(parents=True, exist_ok=True)

    # 1. requirements.json
    reqs_json = []
    for r in BENCHMARK_REQUIREMENTS:
        reqs_json.append({
            "requirement_id": r["requirement_id"],
            "req_code": r["req_code"],
            "title": r["title"],
            "description": r["description"],
            "category": r["category"],
            "severity": r["severity"],
            "expected_status": r["expected_status"],
            "expected_review_state": r["expected_review_state"],
        })
    with open(GT_DIR / "requirements.json", "w", encoding="utf-8") as f:
        json.dump(reqs_json, f, indent=2)

    # 2. evidence_links.json
    links_json = []
    for r in BENCHMARK_REQUIREMENTS:
        for ev in r.get("evidence_sources", []):
            links_json.append({
                "requirement_id": r["requirement_id"],
                "req_code": r["req_code"],
                "document": ev["document"],
                "page": ev["page"],
                "relationship": ev["relationship"],
                "quote": ev["quote"],
            })
    with open(GT_DIR / "evidence_links.json", "w", encoding="utf-8") as f:
        json.dump(links_json, f, indent=2)

    # 3. expected_findings.json
    findings_json = []
    finding_counter = 1
    for r in BENCHMARK_REQUIREMENTS:
        if r["expected_status"] in ("Partial", "Missing", "Conflict"):
            finding_type_map = {
                "Missing": "Missing evidence",
                "Partial": "Partial evidence",
                "Conflict": "Potential conflict",
            }
            findings_json.append({
                "finding_code": f"EXP-F-{finding_counter:03d}",
                "requirement_id": r["requirement_id"],
                "req_code": r["req_code"],
                "expected_status": r["expected_status"],
                "finding_type": finding_type_map.get(r["expected_status"], "Partial evidence"),
                "severity": r["severity"],
                "category": r["category"],
                "expected_evidence": [
                    {"document": ev["document"], "page": ev["page"]}
                    for ev in r.get("evidence_sources", [])
                ],
            })
            finding_counter += 1
    with open(GT_DIR / "expected_findings.json", "w", encoding="utf-8") as f:
        json.dump(findings_json, f, indent=2)

    print(f"Ground truth JSON files exported to {GT_DIR}")


def generate_all():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    GT_DIR.mkdir(parents=True, exist_ok=True)

    print("Generating synthetic technical documents for BCU-800V benchmark...")
    create_product_requirements_srs()
    create_system_architecture_pdf()
    create_oem_supplier_datasheet_pdf()
    create_battery_management_test_report_pdf()
    create_environmental_emc_report_pdf()
    create_thermal_runaway_safety_report_pdf()
    create_compliance_verification_matrix_xlsx()

    print("Exporting ground truth JSON metadata...")
    export_ground_truth_json()
    print("Benchmark dataset generation complete!")


if __name__ == "__main__":
    generate_all()
