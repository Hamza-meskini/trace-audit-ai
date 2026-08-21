"""Generator script for complex, realistic engineering audit test files.

Generates:
1. System_Requirements_Specification_SRS_v2.4.docx (Complex System Spec)
2. Laboratory_Test_Report_TR-2026-894.pdf (Detailed Test Report with measurements)
3. OEM_Supplier_Datasheet_DS-PSU-480W.pdf (Datasheet with subtle contradictions)
4. EMC_Environmental_Compliance_Matrix.xlsx (Excel Matrix with test gaps)
5. Safety_and_Operations_Manual_MAN-042.docx (User & Safety Manual)
"""

import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import fitz  # PyMuPDF


OUT_DIR = Path("c:/Users/Hamza/Desktop/New folder/AudiTrace/sample_documents")
OUT_DIR.mkdir(parents=True, exist_ok=True)


def create_srs_docx():
    """Create a complex 5-page System Requirements Specification (DOCX)."""
    doc = Document()

    # Title & Metadata
    title = doc.add_heading("System Requirements Specification", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Product: Industrial Controller X200 (Series B)\n").bold = True
    p.add_run("Document Ref: SRS-2026-X200-EU | Version: 2.4 | Status: Approved\n")
    p.add_run("Author: Atlas Motion Systems — Systems Engineering Division\n")
    p.add_run("Applicable Directives: 2014/30/EU (EMC), 2014/35/EU (LVD), 2006/42/EC (Machinery)\n")

    doc.add_heading("1. Scope and System Overview", level=1)
    doc.add_paragraph(
        "The Industrial Controller X200 is an embedded, DIN-rail mounted real-time automation controller "
        "designed for high-reliability motion control, robotic cell coordination, and critical industrial telemetry. "
        "This specification establishes mandatory engineering requirements, functional criteria, electrical ratings, "
        "and safety boundaries for CE compliance and industrial deployment."
    )

    doc.add_heading("2. Applicable Harmonized Standards", level=1)
    standards_table = doc.add_table(rows=1, cols=3)
    standards_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = standards_table.rows[0].cells
    hdr[0].text = "Standard Ref"
    hdr[1].text = "Title"
    hdr[2].text = "Compliance Target"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    standards_data = [
        ("IEC 61010-1:2010+A1:2016", "Safety requirements for electrical equipment for measurement & control", "Mandatory (Class I)"),
        ("EN 61326-1:2021", "EMC requirements — Equipment for control and laboratory use (Industrial)", "Zone B Industrial"),
        ("ISO 13849-1:2023", "Safety of machinery — Safety-related parts of control systems", "PL d, Cat 3"),
        ("IEC 60529:2013", "Degrees of protection provided by enclosures (IP Code)", "IP65 Enclosure"),
    ]
    for ref, title_txt, target in standards_data:
        row = standards_table.add_row().cells
        row[0].text = ref
        row[1].text = title_txt
        row[2].text = target

    doc.add_heading("3. Electrical Specifications & Power Budget", level=1)
    
    doc.add_paragraph(
        "REQ-ELE-001 [Critical]: Nominal Operating Voltage and Range\n"
        "The controller shall operate continuously from a regulated DC input supply with a nominal voltage of 24.0 V DC. "
        "The system must remain fully functional across an input tolerance range of 18.0 V DC to 30.0 V DC (24V DC -25% / +25%) "
        "with transient overvoltage withstand up to 36.0 V DC for durations ≤ 500 ms."
    )

    doc.add_paragraph(
        "REQ-ELE-002 [High]: Maximum Power Consumption & Inrush Current\n"
        "Under maximum continuous processing and full digital/analog I/O load (32 channels energized), "
        "total system power consumption shall not exceed 45.0 W. Peak inrush current during cold-start boot sequence "
        "shall be limited to ≤ 8.0 A for a duration not exceeding 10.0 ms."
    )

    doc.add_paragraph(
        "REQ-ELE-003 [Critical]: Galvanic Isolation and Dielectric Withstand\n"
        "All primary field-bus I/O channels, Ethernet interfaces, and internal 3.3V/5V logic power planes "
        "shall feature galvanic optical and magnetic isolation. The insulation barrier shall withstand a test voltage "
        "of 2.5 kV AC RMS (50/60 Hz) for 60 seconds without dielectric breakdown, and maintain leakage current < 1.0 mA."
    )

    doc.add_paragraph(
        "REQ-ELE-004 [High]: Internal DC Bus Ripple and Noise\n"
        "The internal filtered DC power rail distributed to high-speed analog acquisition stages shall exhibit "
        "peak-to-peak ripple and noise not exceeding 50.0 mV pk-pk across a 20 MHz measurement bandwidth under 100% load."
    )

    doc.add_heading("4. Environmental & Climatic Requirements", level=1)

    doc.add_paragraph(
        "REQ-ENV-001 [High]: Operating Ambient Temperature Range\n"
        "The equipment shall maintain full functional performance and timing accuracy across an ambient operating "
        "temperature range of -20.0 °C to +55.0 °C under continuous convection cooling without forced airflow."
    )

    doc.add_paragraph(
        "REQ-ENV-002 [Medium]: Storage and Transport Temperature Limits\n"
        "The equipment, in non-energized packed state, shall withstand storage temperatures from -40.0 °C to +85.0 °C "
        "and relative humidity levels up to 95.0% non-condensing (at +40.0 °C for 96 hours)."
    )

    doc.add_paragraph(
        "REQ-ENV-003 [High]: Ingress Protection (IP Rating)\n"
        "When installed inside a sealed industrial enclosure with properly torqued gland connectors, the front-panel "
        "user interface and sealed bezel shall provide IP65 ingress protection according to IEC 60529 (dust-tight and "
        "protection against water jets from any direction)."
    )

    doc.add_heading("5. Functional Safety & Real-Time Performance", level=1)

    doc.add_paragraph(
        "REQ-SAF-001 [Critical]: Emergency Stop (E-Stop) Response Time\n"
        "Upon actuation of the dual-channel emergency stop circuit (Safe Torque Off - STO), the system shall de-energize "
        "all motion control drive gate signals within a maximum deterministic latency of ≤ 25.0 ms (target ≤ 20.0 ms)."
    )

    doc.add_paragraph(
        "REQ-SAF-002 [Critical]: Dual-Channel Redundancy & Diagnostic Coverage\n"
        "The safety supervisory micro-controller shall achieve diagnostic coverage (DC) ≥ 99.0% and mean time to dangerous "
        "failure (MTTFd) ≥ 100 years, satisfying Category 3 / Performance Level d (PL d) per ISO 13849-1."
    )

    doc.add_paragraph(
        "REQ-COM-001 [Medium]: Real-Time EtherCAT Cycle Jitter\n"
        "The industrial Ethernet interface operating in EtherCAT master mode shall support a deterministic cycle time "
        "of 1.0 ms with maximum allowable cycle jitter ≤ 15.0 microseconds across 1,000,000 consecutive communication frames."
    )

    doc_path = OUT_DIR / "1_System_Requirements_Specification_SRS_v2.4.docx"
    doc.save(str(doc_path))
    print(f"Generated: {doc_path}")


def create_test_report_pdf():
    """Create a complex 4-page Laboratory Test Report PDF with measured data."""
    pdf = fitz.open()

    # Page 1: Cover & Executive Summary
    page1 = pdf.new_page(width=595, height=842) # A4
    page1.insert_text((50, 60), "EURO-TECH LABS — ACCREDITED CONFORMITY TEST REPORT", fontsize=11, fontname="helv", color=(0.2, 0.4, 0.7))
    page1.draw_line((50, 70), (545, 70), color=(0.2, 0.4, 0.7), width=1.5)
    
    page1.insert_text((50, 110), "TEST REPORT: TR-2026-894-ELEC", fontsize=18, fontname="hebo", color=(0.1, 0.1, 0.1))
    page1.insert_text((50, 135), "Product Under Test: Industrial Controller X200 (HW Rev 2.1)", fontsize=12, fontname="hebo")
    page1.insert_text((50, 155), "Manufacturer: Atlas Motion Systems | Test Date: August 12-16, 2026", fontsize=10, fontname="helv")
    page1.insert_text((50, 175), "Accreditation: ISO/IEC 17025 Certified Electrical & Safety Test Lab", fontsize=10, fontname="helv")

    page1.insert_text((50, 220), "1. Executive Summary & Verdict Table", fontsize=14, fontname="hebo")
    summary_text = (
        "This compliance test report details the comprehensive verification of the Industrial Controller X200 "
        "against electrical safety, power boundaries, dielectric insulation, and timing criteria specified in "
        "SRS-2026-X200-EU and harmonized standard IEC 61010-1:2010."
    )
    page1.insert_textbox(fitz.Rect(50, 235, 545, 290), summary_text, fontsize=9.5, fontname="helv")

    # Table on Page 1
    page1.draw_rect(fitz.Rect(50, 300, 545, 450), color=(0.7, 0.7, 0.7), width=1)
    page1.draw_line((50, 325), (545, 325), color=(0.7, 0.7, 0.7), width=1)
    page1.insert_text((60, 317), "Clause / Test Parameter", fontsize=10, fontname="hebo")
    page1.insert_text((260, 317), "Required Spec", fontsize=10, fontname="hebo")
    page1.insert_text((400, 317), "Measured Value", fontsize=10, fontname="hebo")
    page1.insert_text((490, 317), "Verdict", fontsize=10, fontname="hebo")

    tests = [
        ("Input Voltage Operating Range", "18.0 V to 30.0 V DC", "Tested at 18.0V, 24.0V, 30.0V", "PASS"),
        ("Max Power Consumption (Full Load)", "≤ 45.0 W", "38.6 W measured at 24V", "PASS"),
        ("Cold-Start Inrush Current", "≤ 8.0 A, duration ≤ 10 ms", "6.2 A peak for 4.8 ms", "PASS"),
        ("Dielectric Withstand (Chassis/IO)", "2.5 kV AC RMS for 60s", "2.50 kV applied, leakage 0.32 mA", "PASS"),
        ("STO Emergency Stop Latency", "≤ 25.0 ms", "18.4 ms measured", "PASS"),
        ("Thermal Operating Limit Test", "-20.0 °C to +55.0 °C", "Pass at -20°C, Pass at +55°C", "PASS"),
    ]
    y = 345
    for clause, req, measured, verdict in tests:
        page1.insert_text((60, y), clause, fontsize=8.5, fontname="helv")
        page1.insert_text((260, y), req, fontsize=8.5, fontname="helv")
        page1.insert_text((400, y), measured, fontsize=8.5, fontname="helv")
        page1.insert_text((495, y), verdict, fontsize=8.5, fontname="hebo", color=(0.1, 0.6, 0.2))
        y += 20

    # Page 2: Detailed Electrical Measurements
    page2 = pdf.new_page(width=595, height=842)
    page2.insert_text((50, 60), "EURO-TECH LABS | Test Report TR-2026-894-ELEC | Page 2 of 3", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))
    page2.insert_text((50, 100), "2. Electrical Power and Voltage Tolerance Verification (REQ-ELE-001)", fontsize=13, fontname="hebo")
    
    sec2_text = (
        "Test Procedure: The unit under test (EUT) was powered via a programmable DC power supply (Chroma 62050P). "
        "Input voltage was swept from 17.5 V DC to 31.0 V DC in increments of 0.5 V under continuous I/O cycling. "
        "The controller booted and operated without communication dropouts or memory resets from 18.0 V to 30.0 V DC. "
        "Under 24.0 V DC nominal supply, the measured steady-state power draw was 38.6 W with all 32 output channels active. "
        "Transient overvoltage pulse testing: A 36.0 V DC pulse of 500 ms duration was injected at 1-minute intervals. "
        "The internal TVS clamping diodes maintained internal rail stability with zero system resets recorded."
    )
    page2.insert_textbox(fitz.Rect(50, 120, 545, 230), sec2_text, fontsize=9.5, fontname="helv")

    page2.insert_text((50, 250), "3. Dielectric Insulation & High-Pot Testing (REQ-ELE-003)", fontsize=13, fontname="hebo")
    sec3_text = (
        "Test Standard: IEC 61010-1 Clause 6.8 (Dielectric Strength Test).\n"
        "Test Equipment: QuadTech Guardian 6000 HiPot Tester (Cal due: 2027-03-15).\n"
        "Test Points:\n"
        "  - Terminal Block A (24V Power In) to Protective Earth (Chassis Ground)\n"
        "  - Digital I/O Bus Terminals to Isolated MCU Logic Ground Plane\n"
        "Test Voltage: 2,500 V AC RMS (50 Hz), ramp-up time 5.0 s, hold duration 60.0 seconds.\n"
        "Results: Zero breakdown observed. Measured leakage current: 0.32 mA (allowable limit < 1.0 mA). Verdict: PASS."
    )
    page2.insert_textbox(fitz.Rect(50, 270, 545, 410), sec3_text, fontsize=9.5, fontname="helv")

    # Page 3: Functional Safety STO Latency Verification
    page3 = pdf.new_page(width=595, height=842)
    page3.insert_text((50, 60), "EURO-TECH LABS | Test Report TR-2026-894-ELEC | Page 3 of 3", fontsize=9, fontname="helv", color=(0.5, 0.5, 0.5))
    page3.insert_text((50, 100), "4. Functional Safety STO Timing & E-Stop Latency (REQ-SAF-001)", fontsize=13, fontname="hebo")
    sec4_text = (
        "Test Setup: Dual-channel safety relay input STO_A and STO_B were connected to a high-speed logic analyzer "
        "(Keysight Infiniium S-Series, 500 MSa/s). The motor drive PWM output gate signal was monitored across 50 simulated "
        "emergency stop events under full load.\n\n"
        "Timing Results Summary:\n"
        "  - Mean STO Disengage Latency: 18.42 ms\n"
        "  - Maximum Recorded Latency: 21.15 ms (well below 25.0 ms threshold limit)\n"
        "  - Minimum Recorded Latency: 16.80 ms\n"
        "  - Channel Discrepancy Detection: 2.1 ms (switches to Safe State within 5 ms on single-channel failure)\n\n"
        "Verdict: Fully compliant with ISO 13849-1 Category 3 PL d and IEC 61508 SIL 2 requirements."
    )
    page3.insert_textbox(fitz.Rect(50, 120, 545, 300), sec4_text, fontsize=9.5, fontname="helv")

    pdf_path = OUT_DIR / "2_Laboratory_Test_Report_TR-2026-894.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_datasheet_contradiction_pdf():
    """Create a Supplier Component Datasheet PDF containing deliberate parameter contradictions."""
    pdf = fitz.open()

    page = pdf.new_page(width=595, height=842)
    page.insert_text((50, 60), "VOLT-CRAFT POWER TECHNOLOGIES — COMPONENT SPECIFICATION", fontsize=10, fontname="helv", color=(0.8, 0.3, 0.1))
    page.draw_line((50, 70), (545, 70), color=(0.8, 0.3, 0.1), width=1.5)

    page.insert_text((50, 105), "OEM Power Converter Module: VC-PSU-480W", fontsize=16, fontname="hebo", color=(0.1, 0.1, 0.1))
    page.insert_text((50, 125), "Integrated 24V DC Internal Auxiliary Power Supply for Motion Systems", fontsize=10, fontname="helv")

    page.insert_text((50, 160), "1. Technical Specifications Summary", fontsize=12, fontname="hebo")

    # Table containing contradictions with REQ-ENV-001 (+40C vs +55C) and REQ-ELE-004 (120mV ripple vs 50mV)
    page.draw_rect(fitz.Rect(50, 180, 545, 360), color=(0.7, 0.7, 0.7), width=1)
    page.draw_line((50, 205), (545, 205), color=(0.7, 0.7, 0.7), width=1)
    page.insert_text((60, 197), "Parameter", fontsize=10, fontname="hebo")
    page.insert_text((260, 197), "Rated Specification", fontsize=10, fontname="hebo")
    page.insert_text((420, 197), "Notes / Derating", fontsize=10, fontname="hebo")

    params = [
        ("Nominal Output Voltage", "24.0 V DC ±1.0%", "Factory calibrated"),
        ("Maximum Output Current", "20.0 A continuous", "Peak 25.0 A for 3s"),
        ("Output Voltage Ripple & Noise", "120 mV pk-pk maximum", "Measured at 20 MHz [CONTRADICTION]"),
        ("Operating Ambient Temperature", "-10.0 °C to +40.0 °C", "Derate 2.5%/°C above 40°C [CONTRADICTION]"),
        ("Storage Temperature", "-40.0 °C to +85.0 °C", "Standard non-condensing"),
        ("Isolation Voltage (Input/Output)", "3.0 kV AC RMS", "1 minute qualification"),
    ]
    y = 225
    for p_name, p_val, notes in params:
        page.insert_text((60, y), p_name, fontsize=8.5, fontname="helv")
        page.insert_text((260, y), p_val, fontsize=8.5, fontname="hebo" if "CONTRADICTION" in notes else "helv")
        page.insert_text((420, y), notes.replace(" [CONTRADICTION]", ""), fontsize=8.5, fontname="helv")
        y += 22

    page.insert_text((50, 390), "2. Engineering Derating Advisory & Thermal Notes", fontsize=12, fontname="hebo")
    notes_text = (
        "CRITICAL APPLICATION NOTE: The VC-PSU-480W auxiliary power converter is rated for nominal full-load operation "
        "up to +40.0 °C ambient temperature. When operated in enclosed environments exceeding +40.0 °C, the module enters "
        "thermal derating mode (output current limited to 12.0 A at +50.0 °C). Full operation at +55.0 °C is NOT supported "
        "without supplemental forced-air cooling (minimum 2.5 m/s airflow).\n\n"
        "Output Ripple: Under maximum load, internal switching ripple on the secondary rail is guaranteed below 120.0 mV pk-pk. "
        "Applications requiring < 50 mV pk-pk ripple must install an external LC low-pass filter stage."
    )
    page.insert_textbox(fitz.Rect(50, 410, 545, 560), notes_text, fontsize=9.5, fontname="helv")

    pdf_path = OUT_DIR / "3_OEM_Supplier_Datasheet_DS-PSU-480W.pdf"
    pdf.save(str(pdf_path))
    pdf.close()
    print(f"Generated: {pdf_path}")


def create_compliance_matrix_xlsx():
    """Create an EMC & Environmental Verification Matrix in Excel with deliberate missing test gaps."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "EMC_Compliance_Matrix"

    # Header style
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
    data_font = Font(name="Segoe UI", size=10)
    pass_font = Font(name="Segoe UI", size=10, bold=True, color="166534")
    missing_font = Font(name="Segoe UI", size=10, bold=True, color="DC2626")

    headers = [
        "Requirement ID",
        "Test Standard",
        "Test Description",
        "Test Level / Severity",
        "Target Criteria",
        "Measured / Status",
        "Test Lab Certificate",
    ]
    ws.append(headers)

    for col_idx, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    rows = [
        ("REQ-EMC-001", "EN 55032:2015", "Radiated Emissions (30MHz-1GHz)", "Class A Industrial", "Quasi-Peak < 40 dBuV/m", "Pass (Max 36.2 dBuV/m at 240MHz)", "TR-EMC-2026-101"),
        ("REQ-EMC-002", "EN 55032:2015", "Conducted Emissions on AC/DC lines", "Class A Industrial", "Average < 46 dBuV", "Pass (Max 41.0 dBuV at 12.5MHz)", "TR-EMC-2026-101"),
        ("REQ-EMC-003", "EN 61000-4-2", "Electrostatic Discharge (ESD) Immunity", "±6kV Contact / ±8kV Air", "Criterion B (No reset)", "Pass (No errors up to ±8kV Air)", "TR-EMC-2026-102"),
        ("REQ-EMC-004", "EN 61000-4-4", "Electrical Fast Transient (EFT) / Burst", "±2kV Power / ±1kV I/O", "Criterion B (Self-recovering)", "Pass (Tested up to ±2kV 5kHz)", "TR-EMC-2026-102"),
        ("REQ-EMC-005", "EN 61000-4-5", "Surge Immunity (1.2/50us waveform)", "±2kV Line-Earth / ±1kV Line-Line", "Criterion B", "Pass (Surge arrestor clamped at 38V)", "TR-EMC-2026-103"),
        ("REQ-ENV-003", "IEC 60529", "IP65 Water Jet Ingress Test", "6.3mm nozzle @ 12.5 L/min for 3 min", "Zero water penetration inside housing", "NOT TESTED [MISSING EVIDENCE]", "TEST SCHEDULED Q4 2026"),
        ("REQ-COM-001", "IEC 61784-2", "EtherCAT 1ms Jitter Verification", "Cycle Jitter ≤ 15.0 microseconds", "Deterministic timing over 1M frames", "Pass (11.8 microseconds max jitter)", "TR-ETH-2026-009"),
    ]

    for row_idx, row_data in enumerate(rows, 2):
        ws.append(row_data)
        for col_idx in range(1, len(row_data) + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.font = data_font
            if "Pass" in str(cell.value):
                cell.font = pass_font
            elif "NOT TESTED" in str(cell.value):
                cell.font = missing_font

    # Column widths
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 38
    ws.column_dimensions["D"].width = 30
    ws.column_dimensions["E"].width = 32
    ws.column_dimensions["F"].width = 36
    ws.column_dimensions["G"].width = 24

    xlsx_path = OUT_DIR / "4_EMC_Environmental_Compliance_Matrix.xlsx"
    wb.save(str(xlsx_path))
    print(f"Generated: {xlsx_path}")


def create_manual_docx():
    """Create a User and Installation Safety Manual DOCX."""
    doc = Document()
    doc.add_heading("Industrial Controller X200 — User & Installation Manual", level=0)
    
    doc.add_paragraph(
        "Document Code: MAN-X200-EU-042 | Revision: 3.1 | Atlas Motion Systems\n"
        "Target Audience: Qualified Electrical Engineers and Panel Builders\n"
    )

    doc.add_heading("1. Important Safety Warnings and Grounding", level=1)
    doc.add_paragraph(
        "WARNING — ELECTRIC SHOCK HAZARD: Before connecting the 24V DC auxiliary power supply, ensure the "
        "Protective Earth (PE) terminal is securely fastened to the main cabinet ground bus bar with minimum 2.5 mm² copper wire. "
        "The internal isolation barrier is rated for 2,500 V AC dielectric withstand. Under no circumstances should input voltage "
        "exceed 36.0 V DC."
    )

    doc.add_heading("2. Environmental Operating Conditions", level=1)
    doc.add_paragraph(
        "Operating Temperature: The controller is engineered for ambient temperatures from -20.0 °C to +55.0 °C. "
        "Maintain a minimum vertical clearance of 50 mm above and below the DIN rail to allow natural convection airflow. "
        "Relative Humidity: 5% to 95% non-condensing. Storage Temperature: -40.0 °C to +85.0 °C."
    )

    doc.add_heading("3. Emergency Stop (STO) Wiring & Commissioning", level=1)
    doc.add_paragraph(
        "Safe Torque Off (STO) terminals STO1 and STO2 must be wired to dual-channel safety contacts according to "
        "ISO 13849-1 Category 3 architecture. The maximum safety loop response time is guaranteed ≤ 25.0 ms. "
        "Conduct an operational test of the E-Stop circuit every 12 months."
    )

    doc_path = OUT_DIR / "5_Safety_and_Operations_Manual_MAN-042.docx"
    doc.save(str(doc_path))
    print(f"Generated: {doc_path}")


if __name__ == "__main__":
    create_srs_docx()
    create_test_report_pdf()
    create_datasheet_contradiction_pdf()
    create_compliance_matrix_xlsx()
    create_manual_docx()
    print("All 5 complex test documents successfully generated in sample_documents/")
